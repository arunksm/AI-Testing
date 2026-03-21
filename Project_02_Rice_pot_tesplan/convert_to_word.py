import os
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    doc.add_heading('QA Test Plan - VWO Login Dashboard', 0)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Simple Markdown parsing
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            # Basic table detection (ignoring for now as python-docx table creation is complex for direct MD lines)
            # Just adding as plain text for context
            p = doc.add_paragraph(line)
            p.style.font.name = 'Courier New'
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\admin\Documents\GitHub\AI_Testing\AITesterBlueprint2x\Project_02_Rice_pot_tesplan\qa_test_plan.md"
    docx_file = r"C:\Users\admin\Documents\GitHub\AI_Testing\AITesterBlueprint2x\Project_02_Rice_pot_tesplan\QA_Test_Plan_VWO.docx"
    md_to_docx(md_file, docx_file)
