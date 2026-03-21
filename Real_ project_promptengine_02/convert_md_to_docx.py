import os
import sys

# Try to import docx, if not found, we will report it.
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: 'python-docx' library is not installed. Please install it using 'pip install python-docx'.")
    sys.exit(1)

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            # Add some spacing for empty lines
            doc.add_paragraph("")
            continue
        
        # Simple Markdown parsing
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('|'):
            # Very basic table display as fixed-width text
            p = doc.add_paragraph(line)
            p.style.font.name = 'Courier New'
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    # Define paths relative to the script location or absolute
    base_dir = r"c:\Users\admin\Documents\GitHub\AI_Testing\AITesterBlueprint2x\Real_ project_promptengine_02"
    md_file = os.path.join(base_dir, "outputs", "test_plan.md")
    docx_file = os.path.join(base_dir, "outputs", "test_plan_vwo.docx")
    
    convert_md_to_docx(md_file, docx_file)
